"""
Prepara las plantillas .docx oficiales para docxtpl.
SOLO usa etiquetas simples {{ variable }} — SIN loops ni condicionales
dentro del documento, porque Word fragmenta la sintaxis Jinja2 en el XML.

Los datos complejos (ítems, características) se pre-formatean en Python
antes de llamar a doc.render().

Uso: python3 preparar_plantillas.py
"""
import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_text(cell, tag_text):
    """Reemplaza texto de una celda preservando el formato del primer run."""
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = tag_text
            for run in p.runs[1:]:
                run.text = ''
            return
    if cell.paragraphs:
        cell.paragraphs[0].text = tag_text


def set_paragraph_text(paragraph, new_text):
    """Reemplaza texto de un párrafo preservando el formato del primer run."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.text = new_text


def preparar_plantilla_et():
    """
    Convierte TDR Alicate Crimping.docx en especificacion_tecnica_tpl.docx.
    Solo etiquetas simples {{ variable }}, sin loops.
    """
    src = os.path.join(BASE_DIR, 'docs', 'TDR Alicate Crimping.docx')
    dst = os.path.join(BASE_DIR, 'plantillas_docx', 'especificacion_tecnica_tpl.docx')

    if not os.path.exists(src):
        print(f"❌ No se encontró {src}")
        return

    doc = Document(src)

    # ── TABLA 1: Datos del Pedido (5 filas × 2 cols) ──
    t1 = doc.tables[0]
    reemplazos_t1 = {
        0: '{{ centro_costo }}',
        1: '{{ actividad_operativa }}',
        2: '{{ denominacion_adquisicion }}',
        3: '{{ numero_pedido }}',
        4: '{{ meta_anio }}',
    }
    for fila_idx, tag in reemplazos_t1.items():
        set_cell_text(t1.rows[fila_idx].cells[1], tag)
    print("  ✅ Tabla 1 (Datos del Pedido): etiquetas insertadas")

    # ── TABLA 2: Ítems (cabecera + 1 fila de datos) ──
    # En lugar de un loop, ponemos una sola fila con los datos
    # pre-formateados del primer ítem. El sistema generará el docx
    # con todos los ítems ya combinados en un solo texto.
    t2 = doc.tables[1]
    fila_datos = t2.rows[1]
    set_cell_text(fila_datos.cells[0], '{{ items_tabla }}')
    # Vaciar las demás celdas de datos (se fusionará toda la fila)
    for i in range(1, len(fila_datos.cells)):
        set_cell_text(fila_datos.cells[i], '')
    print("  ✅ Tabla 2 (Ítems): etiqueta insertada")

    # ── PÁRRAFOS: Finalidad y Objetivo ──
    for p in doc.paragraphs:
        texto = p.text.strip().lower()
        if 'la finalidad de la presente adquisición' in texto:
            set_paragraph_text(p, '{{ finalidad_publica }}')
            print("  ✅ Finalidad pública: etiqueta insertada")
        elif 'contratar la adquisición de' in texto:
            set_paragraph_text(p, '{{ objetivo }}')
            print("  ✅ Objetivo: etiqueta insertada")

    # ── CARACTERÍSTICAS TÉCNICAS ──
    # Reemplazar el bloque de características con una sola variable
    encontro_sellados = False
    primera_char = True

    for p in doc.paragraphs:
        texto = p.text.strip()

        if 'debidamente sellados' in texto.lower():
            encontro_sellados = True
            continue

        if encontro_sellados and 'reglamentos técnicos' in texto.lower():
            break

        if encontro_sellados and texto:
            if texto.upper() == 'CARACTERISTICAS':
                continue

            if primera_char:
                set_paragraph_text(p, '{{ caracteristicas_texto }}')
                primera_char = False
                print("  ✅ Características: etiqueta insertada")
            else:
                # Vaciar líneas de características sobrantes
                set_paragraph_text(p, '')

    os.makedirs(os.path.dirname(dst), exist_ok=True)
    doc.save(dst)
    print(f"\n📄 Plantilla ET guardada: {dst}")


def preparar_plantilla_ft():
    """
    Convierte FICHA TECNICA - CPU ADVANCE.docx en ficha_tecnica_tpl.docx.
    Solo etiquetas simples {{ variable }}.
    """
    src = os.path.join(BASE_DIR, 'docs',
                       'FICHA TECNICA - UNIDAD CENTRAL DE PROCESAMIENTO_ADVANCE.docx')
    dst = os.path.join(BASE_DIR, 'plantillas_docx', 'ficha_tecnica_tpl.docx')

    if not os.path.exists(src):
        print(f"❌ No se encontró {src}")
        return

    doc = Document(src)

    # ── TABLA 1: N° correlativo y Descripción ──
    t1 = doc.tables[0]
    set_cell_text(t1.rows[1].cells[0], '{{ numero_correlativo }}')
    set_cell_text(t1.rows[1].cells[1], '{{ descripcion_bien }}')
    print("  ✅ Tabla 1 (N° y Descripción)")

    # ── TABLA 2: Bien/Estado/Marca/Modelo/Color/Serie/Características/Carta ──
    t2 = doc.tables[1]
    set_cell_text(t2.rows[1].cells[0], '{{ bien_equipo }}')
    set_cell_text(t2.rows[1].cells[4], '{{ estado_fisico }}')
    set_cell_text(t2.rows[3].cells[0], '{{ marca }}')
    set_cell_text(t2.rows[3].cells[1], '{{ marca }}')
    set_cell_text(t2.rows[3].cells[2], '{{ modelo }}')
    set_cell_text(t2.rows[3].cells[3], '{{ color }}')
    set_cell_text(t2.rows[3].cells[4], '{{ numero_serie }}')
    set_cell_text(t2.rows[5].cells[0], '{{ caracteristicas_texto }}')
    set_cell_text(t2.rows[5].cells[4], '{{ carta_levantamiento }}')
    print("  ✅ Tabla 2 (Bien/Marca/Serie/Características)")

    # ── TABLA 3: Datos Requerimiento ──
    t3 = doc.tables[2]
    set_cell_text(t3.rows[2].cells[0], '{{ pedido_compra }}')
    set_cell_text(t3.rows[2].cells[1], '{{ fecha_pedido }}')
    set_cell_text(t3.rows[2].cells[2], '{{ orden_compra }}')
    set_cell_text(t3.rows[2].cells[3], '{{ fecha_orden }}')
    set_cell_text(t3.rows[4].cells[0], '{{ dependencia }}')
    set_cell_text(t3.rows[4].cells[2], '{{ edificio }}')
    print("  ✅ Tabla 3 (Datos Requerimiento)")

    # ── TABLA 4: Responsable + Proveedor ──
    t4 = doc.tables[3]
    set_cell_text(t4.rows[2].cells[0], '{{ responsable_nombre }}')
    set_cell_text(t4.rows[2].cells[2], '{{ responsable_cargo }}')
    set_cell_text(t4.rows[2].cells[4], '{{ responsable_telefono }}')
    set_cell_text(t4.rows[2].cells[6], '{{ responsable_email }}')
    set_cell_text(t4.rows[4].cells[0], '{{ proveedor_razon_social }}')
    set_cell_text(t4.rows[4].cells[1], '{{ proveedor_direccion }}')
    set_cell_text(t4.rows[4].cells[3], '{{ proveedor_telefono }}')
    set_cell_text(t4.rows[4].cells[5], '{{ orden_compra }}')
    set_cell_text(t4.rows[4].cells[7], '{{ costo }}')
    set_cell_text(t4.rows[6].cells[0], '{{ proveedor_email }}')
    set_cell_text(t4.rows[6].cells[3], '{{ comprobante_guia }}')
    set_cell_text(t4.rows[6].cells[5], '{{ fecha_compra }}')
    set_cell_text(t4.rows[6].cells[7], '{{ garantia }}')
    set_cell_text(t4.rows[8].cells[0], '{{ observaciones }}')
    print("  ✅ Tabla 4 (Responsable/Proveedor)")

    os.makedirs(os.path.dirname(dst), exist_ok=True)
    doc.save(dst)
    print(f"\n📄 Plantilla FT guardada: {dst}")


if __name__ == '__main__':
    print("🏛️  Preparando plantillas .docx con etiquetas simples...\n")

    print("═══ Especificación Técnica (ET) ═══")
    preparar_plantilla_et()

    print("\n═══ Ficha Técnica (FT) ═══")
    preparar_plantilla_ft()

    print("\n" + "=" * 55)
    print("✅ Plantillas listas en plantillas_docx/")
    print("   Solo usan {{ variable }} — sin loops ni lógica")
    print("   Todo el formato oficial se mantiene intacto")
    print("=" * 55)
