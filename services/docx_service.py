"""
Servicio de generación de documentos .docx oficiales.

Estrategia híbrida:
  1. docxtpl renderiza las etiquetas simples {{ variable }} (textos)
  2. python-docx manipula la tabla de ítems fila por fila
     y agrega las características + imagen debajo de la tabla

Esto evita los problemas de fragmentación XML de Jinja2 en Word
y permite construir tablas con el número correcto de filas.
"""
import os
import io
import copy
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from config import Config

# Tamaño estándar para las imágenes referenciales: se fija la ALTURA en
# ALTO_IMAGEN_ESTANDAR y el ANCHO se calcula según la proporción de cada
# imagen, para que ninguna quede distorsionada. Ej.: 15x21 -> 5x7 (misma
# proporción, altura estándar 7).
ALTO_IMAGEN_ESTANDAR = Cm(7)


def _copiar_formato_celda(celda_origen, celda_destino):
    """Copia el formato básico de una celda a otra."""
    for p_dest in celda_destino.paragraphs:
        for p_orig in celda_origen.paragraphs:
            if p_orig.runs:
                p_dest.paragraph_format.alignment = p_orig.paragraph_format.alignment
                break


def _agregar_fila_tabla(tabla, valores, fila_ref_idx=1):
    """
    Agrega una nueva fila a la tabla copiando el formato de la fila de referencia.
    valores: lista de strings, uno por columna.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Copiar la fila de referencia para mantener formato
    fila_ref = tabla.rows[fila_ref_idx]
    tr_ref = fila_ref._tr

    # Crear nueva fila copiando el XML de la referencia
    new_tr = copy.deepcopy(tr_ref)
    tabla._tbl.append(new_tr)

    # Rellenar las celdas con los valores
    nueva_fila = tabla.rows[-1]
    for i, valor in enumerate(valores):
        if i < len(nueva_fila.cells):
            celda = nueva_fila.cells[i]
            for p in celda.paragraphs:
                if p.runs:
                    p.runs[0].text = str(valor)
                    for run in p.runs[1:]:
                        run.text = ''
                else:
                    p.text = str(valor)

    return nueva_fila


def generar_especificacion_tecnica(datos):
    """
    Genera el documento .docx de Especificación Técnica.

    Paso 1: docxtpl renderiza textos simples (tabla superior, finalidad, objetivo)
    Paso 2: python-docx construye la tabla de ítems y las características
    """
    plantilla_path = Config.PLANTILLA_ET

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(
            f"No se encontró la plantilla de ET en: {plantilla_path}. "
            "Coloque el archivo especificacion_tecnica_tpl.docx en plantillas_docx/"
        )

    items = datos.get('items', [])

    # Datos en MAYÚSCULAS para la tabla superior
    datos_render = {
        'centro_costo': (datos.get('centro_costo', '') or '').upper(),
        'actividad_operativa': (datos.get('actividad_operativa', '') or '').upper(),
        'denominacion_adquisicion': (datos.get('denominacion_adquisicion', '') or '').upper(),
        'numero_pedido': datos.get('numero_pedido', ''),
        'meta_anio': (datos.get('meta_anio', '') or '').upper(),
        # Finalidad y objetivo en minúscula (redacción normal)
        'finalidad_publica': datos.get('finalidad_publica', ''),
        'objetivo': datos.get('objetivo', ''),
        # Placeholder para la tabla de ítems (se reemplazará después)
        'items_tabla': '',
        'caracteristicas_texto': '',
    }

    # ── Paso 1: Renderizar textos simples con docxtpl ──
    doc_tpl = DocxTemplate(plantilla_path)
    doc_tpl.render(datos_render)

    # Guardar en buffer temporal
    temp_buffer = io.BytesIO()
    doc_tpl.save(temp_buffer)
    temp_buffer.seek(0)

    # ── Paso 2: Abrir con python-docx y manipular tabla de ítems ──
    doc = Document(temp_buffer)

    # Encontrar la tabla de ítems (Tabla 2: Ítem | DESCRIPCIÓN | CLASIFICADOR | ...)
    tabla_items = None
    for tabla in doc.tables:
        # Buscar la tabla que tiene "Ítem" o "DESCRIPCIÓN" en la cabecera
        primera_fila_textos = [c.text.strip().upper() for c in tabla.rows[0].cells]
        if any('DESCRIPCI' in t for t in primera_fila_textos):
            tabla_items = tabla
            break

    if tabla_items and items:
        # Limpiar la fila de datos existente (fila 1 = placeholder)
        fila_placeholder = tabla_items.rows[1]
        for celda in fila_placeholder.cells:
            for p in celda.paragraphs:
                if p.runs:
                    for run in p.runs:
                        run.text = ''
                else:
                    p.text = ''

        # Escribir el primer ítem en la fila placeholder
        primer_item = items[0]
        valores_primer = [
            str(1).zfill(3),
            (primer_item.get('descripcion', '') or '').upper(),
            primer_item.get('clasificador', '') or '-',
            primer_item.get('unidad_medida', 'UNIDAD'),
            _format_cantidad(primer_item.get('cantidad', 1)),
        ]
        for i, valor in enumerate(valores_primer):
            if i < len(fila_placeholder.cells):
                celda = fila_placeholder.cells[i]
                for p in celda.paragraphs:
                    if p.runs:
                        p.runs[0].text = valor
                        for run in p.runs[1:]:
                            run.text = ''
                    else:
                        p.text = valor

        # Agregar filas adicionales para los demás ítems
        for idx, item in enumerate(items[1:], 2):
            valores = [
                str(idx).zfill(3),
                (item.get('descripcion', '') or '').upper(),
                item.get('clasificador', '') or '-',
                item.get('unidad_medida', 'UNIDAD'),
                _format_cantidad(item.get('cantidad', 1)),
            ]
            _agregar_fila_tabla(tabla_items, valores, fila_ref_idx=1)

    # ── Paso 3: Escribir características e imágenes referenciales intercaladas ──
    # Cada ítem lleva su propio encabezado en negrita "CARACTERÍSTICAS DE
    # <descripción>" y, justo debajo, su imagen referencial (si la tiene),
    # en el orden: características 1 + imagen 1, características 2 + imagen 2...
    _limpiar_imagenes_existentes(doc)
    _escribir_caracteristicas_tituladas(doc, None, None, items=items)

    # Guardar resultado final
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generar_ficha_tecnica(datos):
    """
    Genera el documento .docx de Ficha Técnica rellenando la plantilla oficial.
    Las características se escriben con un encabezado en negrita
    "CARACTERÍSTICAS DE <descripción del bien>" para diferenciar el ítem.
    """
    plantilla_path = Config.PLANTILLA_FICHA

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(
            f"No se encontró la plantilla de FT en: {plantilla_path}. "
            "Coloque el archivo ficha_tecnica_tpl.docx en plantillas_docx/"
        )

    chars = datos.get('caracteristicas', [])
    if not isinstance(chars, list):
        chars = []
    titulo = f"CARACTERÍSTICAS DE {datos.get('bien_descripcion') or 'BIEN'}"

    # Renderizar con un marcador centinela; luego lo reemplazamos con
    # python-docx para poder aplicar negrita al encabezado. El marcador
    # aparece en las 4 columnas del cuadro, pero solo la primera lleva el
    # contenido (las demás se dejan vacías para no duplicar el texto).
    SENTINEL = '__caracteristicas_texto__'
    doc_tpl = DocxTemplate(plantilla_path)
    datos_render = {k: v for k, v in datos.items() if k != 'caracteristicas_texto'}
    datos_render['caracteristicas_texto'] = SENTINEL
    doc_tpl.render(datos_render)

    buffer = io.BytesIO()
    doc_tpl.save(buffer)
    buffer.seek(0)
    doc = Document(buffer)

    ocurrencias = [p for p in _iterar_parrafos_doc(doc) if SENTINEL in p.text]
    for n, p in enumerate(ocurrencias):
        if n == 0:
            _escribir_bloque_caract(p, titulo, chars)
        else:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)

    # Imagen referencial del bien (si la FT la incluye)
    imagen = datos.get('imagen_referencial')
    if imagen:
        _insertar_imagen_doc(doc, imagen)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# Funciones auxiliares
# ============================================================

def _format_cantidad(cant):
    """Formatea cantidad como string decimal."""
    try:
        return f"{float(cant):.2f}"
    except (ValueError, TypeError):
        return str(cant)


def _formatear_todas_caracteristicas(items):
    """Formatea las características de todos los ítems como texto."""
    if not items:
        return ''

    bloques = []
    for i, item in enumerate(items, 1):
        chars = item.get('caracteristicas', [])
        if not chars:
            continue

        lineas = []
        desc = item.get('descripcion', f'Ítem {i}')
        lineas.append(f"CARACTERÍSTICAS - {desc.upper()}:")

        for c in chars:
            nombre = c.get('nombre', '')
            valor = c.get('valor', c.get('valor_sugerido', ''))
            if nombre:
                lineas.append(f"- {nombre}: {valor}")

        bloques.append('\n'.join(lineas))

    return '\n\n'.join(bloques)


def _limpiar_imagenes_existentes(doc):
    """Elimina todas las imágenes incrustadas en párrafos y celdas de tabla."""
    from docx.oxml.ns import qn

    parrafos = list(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                parrafos.extend(celda.paragraphs)

    for p in parrafos:
        for run in p.runs:
            for drawing in run._element.findall(qn('w:drawing')):
                run._element.remove(drawing)
            for pict in run._element.findall(qn('w:pict')):
                run._element.remove(pict)


def _escribir_caracteristicas(doc, items):
    """
    Escribe las características de los ítems en el párrafo bordeado
    (dentro del cuadro) que sigue a 'CARACTERISTICAS'/'SELLADOS'.
    Devuelve el índice de ese párrafo o None.
    """
    target = None
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if '{{ caracteristicas_texto }}' in txt:
            target = p
            target_idx = i
            break
        # Respaldo: párrafo vacío justo después del título CARACTERÍSTICAS
        if txt == '' and i > 0:
            prev = doc.paragraphs[i - 1].text.strip().upper()
            if prev == 'CARACTERISTICAS':
                target = p
                target_idx = i
                break

    if target is None:
        return None

    texto = _formatear_todas_caracteristicas(items)
    if texto:
        if target.runs:
            target.runs[0].text = texto
            for r in target.runs[1:]:
                r.text = ''
        else:
            target.text = texto

    return target_idx


def _encontrar_parrafo_caracteristicas(doc):
    """Devuelve (párrafo, índice) del placeholder de características.

    Busca en párrafos de nivel superior y, si no lo encuentra (la Ficha
    Técnica suele ubicar el placeholder dentro de una celda de tabla),
    también en el interior de las tablas.
    """
    # 1) Párrafos de nivel superior (conserva el índice para la ET)
    for i, p in enumerate(doc.paragraphs):
        if 'caracteristicas_texto' in p.text:
            return p, i
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '' and i > 0:
            prev = doc.paragraphs[i - 1].text.strip().upper()
            # Evita confundir con el título "CARACTERÍSTICAS TÉCNICAS:"
            if 'CARACTER' in prev and 'TÉCNICAS' not in prev and 'TECNICAS' not in prev:
                return p, i

    # 2) Dentro de tablas (Ficha Técnica)
    def _buscar_en_celda(parrafos):
        for j, p in enumerate(parrafos):
            if '{{ caracteristicas_texto }}' in p.text:
                return p
        for j, p in enumerate(parrafos):
            if p.text.strip() == '' and j > 0:
                prev = parrafos[j - 1].text.strip().upper()
                if 'CARACTER' in prev and 'TÉCNICAS' not in prev and 'TECNICAS' not in prev:
                    return p
        return None

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                p = _buscar_en_celda(celda.paragraphs)
                if p is not None:
                    return p, None

    return None, None


def _clonar_parrafo_despues(ref_parrafo):
    """Crea un párrafo vacío (mismo formato/borde que ref) justo después de ref y lo devuelve."""
    from docx.text.paragraph import Paragraph
    new_p = copy.deepcopy(ref_parrafo._p)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_parrafo._p.addnext(new_p)
    return Paragraph(new_p, ref_parrafo._parent)


def _insertar_imagen_en_parrafo(parrafo, imagen):
    """Inserta la imagen referencial en el párrafo indicado. Devuelve True si tuvo éxito."""
    full = os.path.join(Config.BASE_DIR, 'static', imagen)
    if not os.path.exists(full):
        return False
    try:
        run = parrafo.add_run()
        # Altura fija; python-docx calcula el ancho proporcional automáticamente,
        # así la imagen no se distorsiona y no depende de Pillow.
        run.add_picture(full, height=ALTO_IMAGEN_ESTANDAR)
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        return False


def _escribir_caracteristicas_tituladas(doc, titulo, chars, items=None):
    """
    Escribe las características en el cuadro e intercala la imagen referencial
    de cada ítem justo debajo de sus características.

    - Ficha Técnica (items=None): un único bloque cuyo encabezado
      'titulo' se imprime en NEGRITA.
    - Especificación Técnica (items): un encabezado en negrita por ítem
      "CARACTERÍSTICAS DE <descripción>" y, a continuación, su imagen, en el
      orden: características 1 + imagen 1, características 2 + imagen 2, etc.

    Trabaja con referencias a objetos párrafo (no índices) para ser robusto
    tanto si el cuadro está en el cuerpo del documento como dentro de una tabla.
    """
    p, _ = _encontrar_parrafo_caracteristicas(doc)
    if p is None:
        return None

    # Alineación a la izquierda (sin justificar) para evitar espacios
    # irregulares entre palabras en las líneas cortas.
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if items is not None:
        bloques = [
            (f"CARACTERÍSTICAS DE {it.get('descripcion', f'Ítem {n + 1}')}",
             it.get('caracteristicas', []))
            for n, it in enumerate(items)
        ]
        items_list = items
    else:
        bloques = [(titulo, chars if isinstance(chars, list) else [])]
        items_list = None

    actual = p
    for n, (encabezado, lista) in enumerate(bloques):
        # A partir del segundo ítem se usa un párrafo bordeado nuevo insertado
        # justo después del bloque anterior (características o imagen).
        if n > 0:
            actual = _clonar_parrafo_despues(actual)
        if actual is None:
            break

        # Limpiar runs previos del párrafo
        for r in list(actual.runs):
            r._element.getparent().remove(r._element)

        r_tit = actual.add_run(encabezado)
        r_tit.bold = True

        lineas = []
        for c in (lista or []):
            nombre = c.get('nombre', '')
            valor = c.get('valor', c.get('valor_sugerido', ''))
            if nombre:
                lineas.append(f"{nombre}: {valor}")
        cuerpo = '\n'.join(lineas) if lineas else 'Sin características registradas.'

        r_cuer = actual.add_run('\n' + cuerpo)
        r_cuer.bold = False

        # Imagen referencial del ítem, justo después de sus características
        if items_list is not None:
            img = items_list[n].get('imagen_referencial', '')
            if img:
                parrafo_imagen = _clonar_parrafo_despues(actual)
                if parrafo_imagen is None or not _insertar_imagen_en_parrafo(parrafo_imagen, img):
                    # Sin párrafo o falló la imagen: descartar el clon vacío
                    if parrafo_imagen is not None:
                        parrafo_imagen._p.getparent().remove(parrafo_imagen._p)
                else:
                    actual = parrafo_imagen

    return None


def _escribir_bloque_caract(p, titulo, chars):
    """
    Escribe un bloque de características con encabezado en negrita
    (Ficha Técnica, un único bien). Alineado a la izquierda.
    """
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r_tit = p.add_run(titulo)
    r_tit.bold = True

    lineas = []
    for c in (chars or []):
        nombre = c.get('nombre', '')
        valor = c.get('valor', c.get('valor_sugerido', ''))
        if nombre:
            lineas.append(f"{nombre}: {valor}")
    cuerpo = '\n'.join(lineas) if lineas else 'Sin características registradas.'

    r_cuer = p.add_run('\n' + cuerpo)
    r_cuer.bold = False


def _indice_parrafo(doc, p):
    """Devuelve el índice de un párrafo dentro de doc.paragraphs."""
    for i, par in enumerate(doc.paragraphs):
        if par._p is p._p:
            return i
    return -1


def _insertar_imagen_doc(doc, imagen, desde_idx=-1):
    """
    Inserta una imagen referencial (Ficha Técnica) en un párrafo bordeado
    vacío. Devuelve el índice del párrafo donde se insertó, o None.
    """
    p = _siguiente_parrafo_bordeado_vacio(doc, desde_idx)
    if p is None:
        p = _clonar_parrafo_bordeado(doc, desde_idx)
    if p is None:
        return None
    if _insertar_imagen_en_parrafo(p, imagen):
        return _indice_parrafo(doc, p)
    return None


def _tiene_borde(p):
    """Indica si un párrafo tiene borde (está dentro de un cuadro)."""
    pPr = p._p.pPr
    if pPr is None:
        return False
    return pPr.find(qn('w:pBdr')) is not None


def _tiene_imagen(p):
    """Indica si un párrafo ya contiene una imagen embebida."""
    for run in p.runs:
        if run._element.find(qn('w:drawing')) is not None or \
           run._element.find(qn('w:pict')) is not None:
            return True
    return False


def _siguiente_parrafo_bordeado_vacio(doc, desde_idx):
    """Devuelve el próximo párrafo bordeado, vacío y sin imagen."""
    for p in doc.paragraphs[desde_idx + 1:]:
        if _tiene_borde(p) and p.text.strip() == '' and not _tiene_imagen(p):
            return p
    return None


def _clonar_parrafo_bordeado(doc, ref_idx):
    """Crea un párrafo bordeado vacío justo después de ref_idx."""
    from docx.text.paragraph import Paragraph
    ref = doc.paragraphs[ref_idx]
    pPr = ref._p.pPr
    if pPr is None or pPr.find(qn('w:pBdr')) is None:
        return None
    new_p = copy.deepcopy(ref._p)
    # Eliminar el contenido (runs) del clon
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref._p.addnext(new_p)
    return Paragraph(new_p, doc)
